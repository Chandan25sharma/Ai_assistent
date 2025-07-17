# File Processing Tools - PDF, DOCX, Excel, etc.
import os
from pathlib import Path
from typing import Dict, List, Optional, Any
import json

class FileProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.xlsx', '.csv', '.txt', '.json']
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Process file based on its extension"""
        path_obj = Path(file_path)
        
        if not path_obj.exists():
            return {"error": f"File not found: {file_path}"}
        
        extension = path_obj.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._process_pdf(path_obj)
            elif extension == '.docx':
                return self._process_docx(path_obj)
            elif extension in ['.xlsx', '.xls']:
                return self._process_excel(path_obj)
            elif extension == '.csv':
                return self._process_csv(path_obj)
            elif extension == '.txt':
                return self._process_txt(path_obj)
            elif extension == '.json':
                return self._process_json(path_obj)
            else:
                return {"error": f"Unsupported file format: {extension}"}
        
        except Exception as e:
            return {"error": f"Error processing file: {str(e)}"}
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                return {
                    "file_type": "PDF",
                    "file_name": file_path.name,
                    "page_count": len(pdf_reader.pages),
                    "text_content": text_content.strip(),
                    "word_count": len(text_content.split())
                }
        
        except ImportError:
            return {"error": "PyPDF2 not installed. Run: pip install PyPDF2"}
        except Exception as e:
            return {"error": f"PDF processing error: {str(e)}"}
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process DOCX file"""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return {
                "file_type": "DOCX",
                "file_name": file_path.name,
                "paragraph_count": len(doc.paragraphs),
                "text_content": text_content.strip(),
                "word_count": len(text_content.split())
            }
        
        except ImportError:
            return {"error": "python-docx not installed. Run: pip install python-docx"}
        except Exception as e:
            return {"error": f"DOCX processing error: {str(e)}"}
    
    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Process Excel file"""
        try:
            import pandas as pd
            
            # Read all sheets
            sheets_data = pd.read_excel(file_path, sheet_name=None)
            
            result = {
                "file_type": "Excel",
                "file_name": file_path.name,
                "sheet_count": len(sheets_data),
                "sheets": {}
            }
            
            for sheet_name, df in sheets_data.items():
                result["sheets"][sheet_name] = {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist(),
                    "preview": df.head().to_dict('records')
                }
            
            return result
        
        except ImportError:
            return {"error": "pandas/openpyxl not installed. Run: pip install pandas openpyxl"}
        except Exception as e:
            return {"error": f"Excel processing error: {str(e)}"}
    
    def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV file"""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            return {
                "file_type": "CSV",
                "file_name": file_path.name,
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "preview": df.head().to_dict('records'),
                "summary": df.describe().to_dict() if df.select_dtypes(include=['number']).shape[1] > 0 else None
            }
        
        except ImportError:
            return {"error": "pandas not installed. Run: pip install pandas"}
        except Exception as e:
            return {"error": f"CSV processing error: {str(e)}"}
    
    def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """Process text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                "file_type": "Text",
                "file_name": file_path.name,
                "content": content,
                "line_count": len(content.splitlines()),
                "word_count": len(content.split()),
                "char_count": len(content)
            }
        
        except Exception as e:
            return {"error": f"Text processing error: {str(e)}"}
    
    def _process_json(self, file_path: Path) -> Dict[str, Any]:
        """Process JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            return {
                "file_type": "JSON",
                "file_name": file_path.name,
                "content": data,
                "structure": self._analyze_json_structure(data)
            }
        
        except Exception as e:
            return {"error": f"JSON processing error: {str(e)}"}
    
    def _analyze_json_structure(self, data: Any) -> Dict[str, Any]:
        """Analyze JSON structure"""
        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": list(data.keys()),
                "key_count": len(data)
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "item_types": list(set(type(item).__name__ for item in data[:10]))
            }
        else:
            return {
                "type": type(data).__name__,
                "value": str(data)[:100]
            }
    
    def summarize_file(self, file_path: str, ai_manager=None) -> str:
        """Generate AI summary of file content"""
        result = self.process_file(file_path)
        
        if "error" in result:
            return f"❌ {result['error']}"
        
        # Create summary prompt
        content = ""
        if "text_content" in result:
            content = result["text_content"][:2000]  # Limit content length
        elif "content" in result:
            content = str(result["content"])[:2000]
        
        if not ai_manager or not content:
            return f"📄 File processed: {result.get('file_name', 'Unknown')}"
        
        prompt = f"Please provide a concise summary of this {result.get('file_type', 'file')} content:\n\n{content}"
        
        ai_response = ai_manager.get_response(prompt)
        return ai_response.get('response', 'Summary generation failed')
    
    def analyze_file(self, file_path: str, analysis_type: str = "summary") -> Dict[str, Any]:
        """Analyze file content based on analysis type"""
        try:
            # First process the file to get its content
            file_data = self.process_file(file_path)
            
            if "error" in file_data:
                return file_data
            
            content = file_data.get("content", "")
            
            # Basic analysis
            analysis = {
                "file_path": file_path,
                "analysis_type": analysis_type,
                "file_size": os.path.getsize(file_path),
                "file_type": Path(file_path).suffix,
                "content_length": len(content),
                "word_count": len(content.split()) if content else 0,
                "line_count": content.count('\n') if content else 0
            }
            
            if analysis_type == "summary":
                # Create a simple summary
                lines = content.split('\n')[:10]  # First 10 lines
                analysis["summary"] = '\n'.join(lines)
            
            elif analysis_type == "keywords":
                # Extract keywords (simple approach)
                words = content.lower().split()
                word_freq = {}
                for word in words:
                    if len(word) > 3:  # Only words longer than 3 chars
                        word_freq[word] = word_freq.get(word, 0) + 1
                
                # Top 10 most frequent words
                analysis["keywords"] = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
            
            elif analysis_type == "structure":
                # Analyze structure
                analysis["structure"] = {
                    "paragraphs": content.count('\n\n'),
                    "sentences": content.count('.') + content.count('!') + content.count('?'),
                    "characters": len(content),
                    "characters_no_spaces": len(content.replace(' ', ''))
                }
            
            return analysis
            
        except Exception as e:
            return {"error": f"Analysis failed: {str(e)}"}

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return self.supported_formats.copy()
